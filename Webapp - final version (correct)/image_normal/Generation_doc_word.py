from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from PIL import Image as PILIMAGE
from io import BytesIO
from datetime import date
from flask import flash,request,session
import pyperclip
from datetime import datetime
from database_connection import execute_query,db_config_client
import math
import mysql.connector

current_date = date.today()
def generate_report_document(session_folder_path):
    email = session.get('email')
    date = current_date
    name = request.cookies.get('names')
    feeder = request.cookies.get('feeder')
    zone = request.cookies.get('zone')
    groupement = request.cookies.get('groupement')
    # Fonction pour créer une table des matières automatique
    def create_table_of_contents(document):
        #////// Insérer la table des matières
        table_of_contents = document.add_paragraph("Table des matières")
        table_of_contents.runs[0].font.size = Pt(16)
        table_of_contents.runs[0].font.name = "Times New Roman"
        table_of_contents.runs[0].underline = True
        table_of_contents.runs[0].font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
        document.add_paragraph(f"RAPPORT D’INSPECTION PAR DRONE DANS LA ZONE {zone}")
        feeder_title = document.add_paragraph()
        run = feeder_title.add_run("FEEDER : ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        document.add_paragraph(f"\t{feeder}")
        groupement_title = document.add_paragraph()
        run = groupement_title.add_run("GROUPEMENT : ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        document.add_paragraph(f"\tGROUPEMENT TRONCONS ENTRE {groupement}")
        return document
    # Créer le document Word
    document = Document()
    # Paramètres de mise en page
    sections = document.sections
    for section in sections:
        section.top_margin = Pt(52)  # Marge supérieure
        section.bottom_margin = Pt(52)  # Marge inférieure
        section.left_margin = Pt(52)  # Marge gauche
        section.right_margin = Pt(52)  # Marge droite
        section.page_width = Pt(612)  # Largeur de page
        section.page_height = Pt(792)  # Hauteur de page

    # Ajouter la première page
    first_page = document.add_paragraph()
    first_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = first_page.add_run(f"RAPPORT D’INSPECTION PAR DRONE DANS LA ZONE {zone}")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)

    # Créer la table des matières
    document = create_table_of_contents(document)
    # Ajouter la distance entre l'entête et la page
    for section in sections:
        section.header_distance = Pt(36)

    # Ajouter la deuxième page
    document.add_page_break()
    second_page = document.add_paragraph()
    second_page.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = second_page.add_run("FEEDER:")
    run.bold = True
    run.underline = True
    run.font.size = Pt(20)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
    second_page.add_run(f"\n\n\t{feeder}\n\n")
    run = second_page.add_run("GROUPEMENT:")
    run.bold = True
    run.underline = True
    run.font.size = Pt(20)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Bleu sombre (#2F5496)
    second_page.add_run(f"\n\n\tGROUPEMENT TRONCONS ENTRE {groupement}")
    run.font.size = Pt(20)
    run.font.name = "Times New Roman"
    # Ajouter la troisième page
    document.add_page_break()
    third_page = document.add_paragraph()
    third_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = third_page.add_run("\n")
    run.add_picture("./logo/logo.png", width=Inches(7))
    run = third_page.add_run("\n\n")
    run = third_page.add_run(f"GROUPEMENT TRONCONS ENTRE {groupement}")
    run.underline = True
    run.font.size = Pt(16)
#Parcours des images dans le dossier de session
    for root, dirs, files in os.walk(session_folder_path):
        for image_file in files:
            if image_file.lower().endswith(('.jpg', '.png', '.jpeg')):
                image_path = os.path.join(root, image_file)
                image = PILIMAGE.open(image_path)  # Charger l'image à partir du chemin
                # Ajoutez ici la logique pour extraire les informations de l'image, par exemple :
                parent_dir = os.path.basename(os.path.dirname(image_path))
                defect = f"{parent_dir}"
                
                # Récupérer le nom du fichier KML correspondant à l'image
                kml_filename = os.path.splitext(image_file)[0] + ".kml"
                kml_path = os.path.join(session_folder_path, "Data_kml", kml_filename)
                
                # Ajouter la quatrième page
                document.add_page_break()
                fourth_page = document.add_paragraph()
                fourth_page.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # Créer le tableau pour stocker les données de l'image
                table = document.add_table(rows=5, cols=1)  # Suppression de la ligne pour le fichier KML
                table.style = "Table Grid"
                table.columns[0].width = Pt(900)
                # Cellule pour les défauts
                cell = table.cell(0, 0)
                cell.text = f"Défauts : {defect}"
                cell.paragraphs[0].runs[0].bold = True
                # Reduce image dimensions
                max_image_size = (800, 600)
                try:
                    image.thumbnail(max_image_size, PILIMAGE.ANTIALIAS)
                except:
                    image.thumbnail(max_image_size)
                # Compress image to reduce size
                image = image.convert("RGB")
                image_bytes = BytesIO()
                image.save(image_bytes, format='JPEG', quality=95)
                image_bytes.seek(0)
                # Cellule pour l'image
                cell = table.cell(1, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                cell.add_paragraph().add_run().add_picture(image_bytes, width=Inches(7))
                # Cellule pour la localisation
                cell = table.cell(2, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                cell.add_paragraph("Veuillez télécharger le fichier KML correspond depuis votre application")
                                # Vérifier si le fichier KML existe
                if os.path.exists(kml_path):
                    # Ajouter le nom du fichier KML sous forme de lien
                    cell.add_paragraph().add_run(kml_filename)
                    # Copiez le chemin complet du fichier KML dans le presse-papiers
                    pyperclip.copy(kml_path)
                # Cellule pour la remarque
                cell = table.cell(3, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                cell.add_paragraph().add_run("Remarque: ").bold = True
                cell.add_paragraph("Nous avons constaté une corrosion au niveau du support.")
                # Cellule pour le conseil
                cell = table.cell(4, 0)
                cell.vertical_alignment = 1  # Alignement vertical centré
                cell.add_paragraph().add_run("Conseil: ").bold = True
                cell.add_paragraph("Veuillez effectuer un remplacement du support ou mettre une peinture anti-corrosion")
                
    # Ajouter le pied de page
    footer_text = f"RAPPORT DU {date}\t\t{name}"
    sections[-1].footer.paragraphs[0].text = footer_text
    # En-tête
    header_text = "VOCASARA S.U.A.R.L"
    header_paragraph = sections[0].header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_run = header_paragraph.add_run(header_text)
    header_run.bold = True
    header_run.font.size = Pt(22)
    header_run.font.name = "Times New Roman"
    header_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Noir
    header_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Logo
    logo_path = "./logo/logo.png"
    header_paragraph.add_run().add_picture(logo_path, width=Inches(4)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Sauvegarder le document Word
    flash('document rapport généré avec succés')
    date = datetime.now().strftime("%Y-%m-%d")  # Convert the date to a string with the format "YYYY-MM-DD"
    #document.save("rapport_du_feeder_{}_du_{}.docx".format(feeder,date))
    rapport_folder_path = os.path.join(session_folder_path, "Rapport_defauts_visible")
    if not os.path.exists(rapport_folder_path):
        os.makedirs(rapport_folder_path)
    nom_du_fichier = "Rapport_du_feeder_{}_du_{}.docx".format(feeder,date)
    document_path = os.path.join(rapport_folder_path,nom_du_fichier)
    document.save(document_path)
    file_size_bytes = os.path.getsize(document_path)
    file_size = convert_size(file_size_bytes)
    file_type = 'word/docx'
    insert_file_info_to_db(email, nom_du_fichier,document_path,file_type,date,file_size)

def insert_file_info_to_db(email, file_name, file_path, file_type, creation_date, file_size):
    conn = mysql.connector.connect(**db_config_client)
    cursor = conn.cursor()
    query = "INSERT INTO fichiers (email_utilisateur, nom_fichier, chemin_fichier, type_fichier, date_creation, poids_fichier) " \
            "VALUES (%s, %s, %s, %s, %s, %s)"
    values = (email, file_name, file_path, file_type, creation_date, file_size)
    cursor.execute(query, values)
    conn.commit()
    cursor.close()
    conn.close()

def convert_size(size_bytes):
    if size_bytes == 0:
        return "0B"
    size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
    i = 0 if size_bytes == 0 else int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return "%s %s" % (s, size_name[i])